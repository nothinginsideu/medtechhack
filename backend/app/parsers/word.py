import docx
import typing
from app.parsers.base import BaseParser, ParsedItem, detect_currency, parse_price
from decimal import InvalidOperation
import re

class WordParser(BaseParser):
    def parse(self) -> typing.List[ParsedItem]:
        doc = docx.Document(self.file_path)
        items = []
        raw_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        
        file_currency = detect_currency("", self.file_path)
        
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                # Пропускаем шапку (грубо)
                if i == 0 and self.config.get("has_header", True):
                    continue
                
                cells = [cell.text.strip() for cell in row.cells]
                raw_parts.append(" | ".join(cells))
                if len(cells) >= 2:
                    price_positions = []
                    for idx, cell in enumerate(cells):
                        currency = detect_currency(cell)
                        parsed = parse_price(cell, currency if currency != "KZT" else file_currency)
                        if parsed is not None:
                            price_positions.append((idx, parsed))
                    
                    if price_positions:
                        try:
                            price_idx, price_val = price_positions[0]
                            nonresident_val = price_positions[1][1] if len(price_positions) > 1 else None
                            name_cells = [cell for idx, cell in enumerate(cells) if idx not in {pos[0] for pos in price_positions}]
                            name_cell = " ".join(name_cells).strip()
                            name_cell = re.sub(r'^\d+[\s.)]*', '', name_cell).strip()
                            if len(name_cell) < 5:
                                continue

                            cell_currency = detect_currency(cells[price_idx])
                            currency = cell_currency if cell_currency != "KZT" else file_currency
                                
                            item = ParsedItem(
                                service_name_raw=name_cell,
                                price_resident_kzt=price_val,
                                price_nonresident_kzt=nonresident_val,
                                currency_original=currency
                            )
                            item.validate_prices()
                            items.append(item)
                        except (ValueError, InvalidOperation):
                            continue

        self.raw_content = "\n".join(raw_parts)
        return items
